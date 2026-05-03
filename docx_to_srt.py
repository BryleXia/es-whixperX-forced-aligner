"""
通用脚本：将日语学院提供的双列表格 docx 转换为占位时间戳 SRT。

用法：
  python docx_to_srt.py --input-dir ./日语/muse-raw20.21 --output-dir ./日语/muse-raw20.21
"""

import argparse
import re
import unicodedata
from difflib import SequenceMatcher
from pathlib import Path


def _normalize(text: str) -> str:
    """统一空白字符并去除首尾空白。"""
    text = text.replace("\r", "\n")
    # 合并多个空行为一个
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _split_paragraphs(cell_text: str):
    """按换行拆分为段落，过滤空行。"""
    text = _normalize(cell_text)
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    return paragraphs


def _seconds_to_srt_time(s: float) -> str:
    """秒 → SRT 时间戳格式 HH:MM:SS,mmm"""
    ms = int(round((s % 1) * 1000))
    total_s = int(s)
    hh = total_s // 3600
    mm = (total_s % 3600) // 60
    ss = total_s % 60
    return f"{hh:02d}:{mm:02d}:{ss:02d},{ms:03d}"


def _build_srt(subtitles: list[str], line_duration: float = 10.0) -> str:
    """
    把字幕列表转为 SRT 格式（占位时间戳）。
    """
    lines = []
    cursor = 0.0
    for i, text in enumerate(subtitles, 1):
        start = cursor
        end = cursor + line_duration
        lines.append(str(i))
        lines.append(f"{_seconds_to_srt_time(start)} --> {_seconds_to_srt_time(end)}")
        lines.append(text)
        lines.append("")
        cursor = end
    return "\n".join(lines)


def _find_best_audio_match(docx_name_no_zh: str, audio_stems: list[str]) -> str | None:
    """
    用模糊匹配找到最相近的音频 stem。
    docx_name_no_zh: 如 'ja_tour_muse_0020_seg001'
    audio_stems: 如 ['ja_tour_muse_0020_seg001_tgt', 'ja_tour_muse_0020_seg002_tgt']
    """
    if docx_name_no_zh in audio_stems:
        return docx_name_no_zh

    best = None
    best_score = 0.0
    for stem in audio_stems:
        score = SequenceMatcher(None, docx_name_no_zh, stem).ratio()
        if score > best_score:
            best_score = score
            best = stem

    # 阈值 0.6 过滤掉完全不相关的匹配
    if best_score >= 0.6:
        return best
    return None


def process_docx(docx_path: Path, audio_stems: list[str], default_suffix: str = "_tgt") -> tuple[str, list[str]] | None:
    """
    处理单个 docx，返回 (srt_stem, subtitles) 或 None。
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("请先安装 python-docx：pip install python-docx")

    doc = Document(str(docx_path))
    if not doc.tables:
        print(f"  [跳过] {docx_path.name}：没有表格")
        return None

    table = doc.tables[0]
    if len(table.rows) < 2:
        print(f"  [跳过] {docx_path.name}：表格行数不足")
        return None

    # 收集第二列（跳过表头）
    subtitles = []
    header_skipped = False
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        if not header_skipped:
            header_skipped = True
            continue
        cell_text = cells[1].text
        subtitles.extend(_split_paragraphs(cell_text))

    if not subtitles:
        print(f"  [跳过] {docx_path.name}：未提取到任何字幕文本")
        return None

    # 文件名映射：去掉 zh- 前缀
    raw_name = docx_path.stem  # 如 'zh-ja_tour_muse_0020_seg001'
    if raw_name.startswith("zh-"):
        candidate = raw_name[3:]  # 'ja_tour_muse_0020_seg001'
    else:
        candidate = raw_name

    matched_stem = _find_best_audio_match(candidate, audio_stems)
    if matched_stem:
        srt_stem = matched_stem
    else:
        srt_stem = candidate + default_suffix
        print(f"  [提示] {docx_path.name} 未找到匹配音频，使用默认命名: {srt_stem}")

    return srt_stem, subtitles


def main():
    parser = argparse.ArgumentParser(description="docx 双列表格 → SRT 占位时间戳")
    parser.add_argument("--input-dir", required=True, help="包含 docx 的目录")
    parser.add_argument("--output-dir", required=True, help="SRT 输出目录")
    parser.add_argument("--audio-dir", default=None, help="音频目录（默认与 input-dir 相同）")
    parser.add_argument("--line-duration", type=float, default=10.0, help="每行占位时长（秒），默认 10")
    parser.add_argument("--default-suffix", default="_tgt", help="无匹配音频时的默认后缀")
    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    output_dir = Path(args.output_dir)
    audio_dir = Path(args.audio_dir) if args.audio_dir else input_dir
    output_dir.mkdir(parents=True, exist_ok=True)

    # 扫描音频获取所有 stem（用于模糊匹配）
    audio_files = sorted(
        f for ext in ("*.m4a", "*.mp3", "*.wav", "*.flac")
        for f in audio_dir.glob(ext)
    )
    audio_stems = [f.stem for f in audio_files]
    print(f"发现 {len(audio_stems)} 个音频文件")
    if audio_stems:
        print(f"  音频 stems: {', '.join(audio_stems)}")

    # 扫描 docx
    docx_files = sorted(input_dir.glob("*.docx"))
    print(f"发现 {len(docx_files)} 个 docx 文件")

    converted = 0
    for docx_path in docx_files:
        result = process_docx(docx_path, audio_stems, default_suffix=args.default_suffix)
        if result is None:
            continue
        srt_stem, subtitles = result
        srt_path = output_dir / f"{srt_stem}.asr.qc.srt"
        srt_content = _build_srt(subtitles, line_duration=args.line_duration)
        srt_path.write_text(srt_content, encoding="utf-8")
        print(f"  [OK] {docx_path.name} → {srt_path.name} ({len(subtitles)} 条字幕)")
        converted += 1

    print(f"\n共转换 {converted}/{len(docx_files)} 个 docx → SRT")
    print(f"输出目录: {output_dir}")


if __name__ == "__main__":
    main()
